import os
import json
import hashlib
import requests
from typing import List
from docx import Document as DocxDocument
from langchain_postgres import PGVector
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.abspath(os.path.join(CURRENT_DIR, "..", "..", ".."))



# 配置部分 
COLLECTION_NAME = "parent_child_db_1024"
# DB_CONNECTION = os.getenv("PGVECTOR_CONNECTION", "postgresql+psycopg2://postgres:postgres@192.168.100.160:5432/vector_db")
DB_CONNECTION = os.getenv("PGVECTOR_CONNECTION", "postgresql+psycopg2://hbch:hbch2711@10.32.10.160:5435/postgres")

LOCAL_API_URL = "http://10.32.10.160:8991/embed" # 你的本地接口地址

# 图片保存的本地相对路径
IMAGE_SAVE_DIR = os.path.join(BASE_DIR, "static", "images")

# 图片访问
IMAGE_BASE_URL = "http://localhost:8001/static/images/"
# 文件路径
folder_path = os.path.join(BASE_DIR, "data")

# 加载区划映射表
MAPPING_PATH = os.path.join(CURRENT_DIR, "region_mapping.json")
region_mapping = {}
if os.path.exists(MAPPING_PATH):
    with open(MAPPING_PATH, "r", encoding="utf-8") as f:
        region_mapping = json.load(f)
    print(f"📋 已加载区划映射: {len(region_mapping)} 个文件")
else:
    print("⚠️ 未找到 region_mapping.json，入库时 region_code 将为默认值")

# 自定义本地模型接口类 
class CustomLocalEmbeddings(Embeddings):
    def __init__(self, api_url: str):
        self.api_url = api_url

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """调用本地接口批量生成向量"""
        try:
            response = requests.post(
                self.api_url,
                json=texts,
                headers={"Content-Type": "application/json", "accept": "application/json"},
                timeout=60 # 灌库时可能一次性传很多文本，稍微把超时时间设长一点
            )
            response.raise_for_status()
            # 从返回结果的字典中提取二维数组
            return response.json()["embeddings"]
        except Exception as e:
            print(f"❌ 请求 Embedding 接口失败: {e}")
            return []

    def embed_query(self, text: str) -> List[float]:
        """调用本地接口生成单条文本的向量"""
        embeddings = self.embed_documents([text])
        if embeddings and len(embeddings) > 0:
            return embeddings[0]
        return []


# 初始化模型 
print(f"🔄 加载自定义接口模型: {LOCAL_API_URL} ...")
EMBEDDING_MODEL = CustomLocalEmbeddings(api_url=LOCAL_API_URL)

# 确保图片目录存在
if not os.path.exists(IMAGE_SAVE_DIR):
    os.makedirs(IMAGE_SAVE_DIR)

# 图片提取 
def extract_images_from_paragraph(paragraph):
    """
    检查 Word 段落中是否有图片，如果有，保存图片并返回 Markdown 链接。
    """
    md_links = []
    
    for run in paragraph.runs:
        if 'drawing' in run._element.xml:
            blips = run._element.xpath('.//a:blip')
            for blip in blips:
                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id:
                    image_part = paragraph.part.related_parts[embed_id]
                    image_bytes = image_part.blob
                    
                    img_hash = hashlib.md5(image_bytes).hexdigest()
                    content_type = image_part.content_type
                    ext = content_type.split('/')[-1] if '/' in content_type else 'png'
                    if ext == 'jpeg': ext = 'jpg'
                    
                    filename = f"{img_hash}.{ext}"
                    save_path = os.path.join(IMAGE_SAVE_DIR, filename)
                    
                    if not os.path.exists(save_path):
                        with open(save_path, "wb") as f:
                            f.write(image_bytes)
                    
                    img_url = f"{IMAGE_BASE_URL}{filename}"
                    md_links.append(f"\n![image]({img_url})\n")
    
    return "".join(md_links)


def read_and_process_word(file_path):
    """
    解析 Word：
    """
    try:
        doc = DocxDocument(file_path)
        file_name = os.path.basename(file_path)
        
        full_text_list = []
        for para in doc.paragraphs:
            text = para.text.strip()
            img_md = extract_images_from_paragraph(para)
            
            if text:
                full_text_list.append(text)
            if img_md:
                full_text_list.append(img_md)
        
        full_text_with_images = "\n".join(full_text_list)
        docs_to_save = []
        
        parent_blocks = full_text_with_images.split("***")
        
        for p_idx, parent_block in enumerate(parent_blocks):
            if not parent_block.strip(): continue
            
            parent_title = ""
            lines = parent_block.strip().split('\n')
            for line in lines:
                if "文档名：" in line:
                    parent_title = line.strip()
                    break
            if not parent_title:
                parent_title = f"文档名：{file_name}"
            
            full_parent_context = parent_block.strip()
            child_blocks = parent_block.split("<--split-->")
            
            for c_idx, child_content in enumerate(child_blocks):
                child_content = child_content.strip()
                if not child_content: continue
                
                content_to_vectorize = f"{parent_title}\n{child_content}"
                
                metadata = {
                    "source": file_name,
                    "type": "child_chunk",
                    "recall_context": full_parent_context,
                    "chunk_id": f"{file_name}_p{p_idx}_c{c_idx}",
                    "region_code": region_mapping.get(file_name, "000000")
                }
                
                new_doc = Document(page_content=content_to_vectorize, metadata=metadata)
                docs_to_save.append(new_doc)

        return docs_to_save

    except Exception as e:
        print(f" 解析文件 {file_path} 失败: {e}")
        import traceback
        traceback.print_exc()
        return []


def ingest():
    print(f"连接 PostgreSQL 数据库并初始化集合: {COLLECTION_NAME} (维度: 1024)")
    vector_store = PGVector(
        embeddings=EMBEDDING_MODEL,
        collection_name=COLLECTION_NAME,
        connection=DB_CONNECTION,
        use_jsonb=True,
    )

    if not os.path.exists(folder_path):
        print("文件夹不存在")
        return

    # 清空整个集合，避免重复数据
    print(f"清空集合: {COLLECTION_NAME} ...")
    vector_store.delete_collection()
    vector_store.create_collection()

    all_docs = []
    file_list = [f for f in os.listdir(folder_path) if f.endswith(".docx")]

    print(f"📂 开始处理 {len(file_list)} 个文档...")
    for f in file_list:
        path = os.path.join(folder_path, f)
        docs = read_and_process_word(path)
        all_docs.extend(docs)
        print(f"  - {f}: 生成 {len(docs)} 个切片")

    if all_docs:
        print(f"正在写入 {len(all_docs)} 个切片到PostgreSQL数据库...")
        vector_store.add_documents(documents=all_docs, batch_size=32)
        print("✅ 入库完成！")
    else:
        print("无数据入库")

if __name__ == "__main__":
    ingest()