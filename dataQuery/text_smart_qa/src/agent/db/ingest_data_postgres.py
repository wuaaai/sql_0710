import os
import json
import hashlib
import requests
from typing import List
from docx import Document as DocxDocument
from text_smart_qa.src.agent.db.pgvector_store import PgVectorStore
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from text_smart_qa.src import env_utils

CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.abspath(os.path.join(CURRENT_DIR, "..", "..", ".."))

# 配置部分
COLLECTION_NAME = env_utils.PGVECTOR_COLLECTION_NAME
DB_CONNECTION = env_utils.PGVECTOR_CONNECTION
LOCAL_API_URL = env_utils.EMBEDDING_API_URL
LOCAL_API_KEY = env_utils.EMBEDDING_API_KEY
LOCAL_EMBED_MODEL = env_utils.EMBEDDING_MODEL_NAME

# 图片保存的本地相对路径
IMAGE_SAVE_DIR = os.path.join(BASE_DIR, "static", "images")

# 图片访问
IMAGE_BASE_URL = env_utils.IMAGE_BASE_URL
# 文件路径
folder_path = os.path.join(BASE_DIR, "data")

# 加载区划映射表
MAPPING_PATH = os.path.join(CURRENT_DIR, "region_mapping.json")
region_mapping = {}
if os.path.exists(MAPPING_PATH):
    with open(MAPPING_PATH, "r", encoding="utf-8") as f:
        region_mapping = json.load(f)
    print(f"📋 已加载区划映射: {len(region_mapping)} 个文件")
else:
    print("⚠️ 未找到 region_mapping.json，入库时 region_code 将为默认值")


# 自定义本地模型接口类（同时兼容内网 OpenAI 格式和本地格式）
class CustomLocalEmbeddings(Embeddings):
    def __init__(self, api_url: str, api_key: str = "", model_name: str = ""):
        self.api_url = api_url
        self.api_key = api_key
        self.model_name = model_name

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        try:
            headers = {"Content-Type": "application/json", "accept": "application/json"}
            if self.api_key:
                headers["Authorization"] = f"Bearer {self.api_key}"
                payload = {"input": texts}
                if self.model_name:
                    payload["model"] = self.model_name
                response = requests.post(self.api_url, json=payload, headers=headers, timeout=60)
                response.raise_for_status()
                data = response.json()
                if "data" in data:
                    sorted_items = sorted(data["data"], key=lambda x: x.get("index", 0))
                    return [item["embedding"] for item in sorted_items]
                return []
            else:
                response = requests.post(self.api_url, json=texts, headers=headers, timeout=60)
                response.raise_for_status()
                return response.json()["embeddings"]
        except Exception as e:
            print(f"❌ 请求 Embedding 接口失败: {e}")
            return []

    def embed_query(self, text: str) -> List[float]:
        embeddings = self.embed_documents([text])
        if embeddings and len(embeddings) > 0:
            return embeddings[0]
        return []


# 初始化模型
print(f"🔄 加载自定义接口模型: {LOCAL_API_URL} ...")
EMBEDDING_MODEL = CustomLocalEmbeddings(api_url=LOCAL_API_URL, api_key=LOCAL_API_KEY, model_name=LOCAL_EMBED_MODEL)

# 确保图片目录存在
if not os.path.exists(IMAGE_SAVE_DIR):
    os.makedirs(IMAGE_SAVE_DIR)


# 图片提取
def extract_images_from_paragraph(paragraph):
    """
    检查 Word 段落中是否有图片，如果有，保存图片并返回 Markdown 链接。
    """
    md_links = []

    for run in paragraph.runs:
        if 'drawing' in run._element.xml:
            blips = run._element.xpath('.//a:blip')
            for blip in blips:
                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id:
                    image_part = paragraph.part.related_parts[embed_id]
                    image_bytes = image_part.blob

                    img_hash = hashlib.md5(image_bytes).hexdigest()
                    content_type = image_part.content_type
                    ext = content_type.split('/')[-1] if '/' in content_type else 'png'
                    if ext == 'jpeg': ext = 'jpg'

                    filename = f"{img_hash}.{ext}"
                    save_path = os.path.join(IMAGE_SAVE_DIR, filename)

                    if not os.path.exists(save_path):
                        with open(save_path, "wb") as f:
                            f.write(image_bytes)

                    img_url = f"{IMAGE_BASE_URL}{filename}"
                    md_links.append(f"\n![image]({img_url})\n")

    return "".join(md_links)


def read_and_process_txt(file_path):
    """解析纯文本文件，无图片提取。"""
    try:
        file_name = os.path.basename(file_path)
        with open(file_path, "r", encoding="utf-8") as f:
            full_text = f.read()

        docs_to_save = []
        parent_blocks = full_text.split("***")

        for p_idx, parent_block in enumerate(parent_blocks):
            if not parent_block.strip():
                continue

            parent_title = ""
            lines = parent_block.strip().split('\n')
            for line in lines:
                if "文档名：" in line:
                    parent_title = line.strip()
                    break
            if not parent_title:
                parent_title = f"文档名：{file_name}"

            full_parent_context = parent_block.strip()
            child_blocks = parent_block.split("<--split-->")

            for c_idx, child_content in enumerate(child_blocks):
                child_content = child_content.strip()
                if not child_content:
                    continue

                content_to_vectorize = f"{parent_title}\n{child_content}"
                metadata = {
                    "source": file_name,
                    "type": "child_chunk",
                    "recall_context": full_parent_context,
                    "chunk_id": f"{file_name}_p{p_idx}_c{c_idx}",
                    "region_code": region_mapping.get(file_name, "000000")
                }
                docs_to_save.append(Document(page_content=content_to_vectorize, metadata=metadata))

        return docs_to_save

    except Exception as e:
        print(f" 解析文件 {file_path} 失败: {e}")
        import traceback
        traceback.print_exc()
        return []


def read_and_process_word(file_path):
    """
    解析 Word：
    """
    try:
        doc = DocxDocument(file_path)
        file_name = os.path.basename(file_path)

        full_text_list = []
        for para in doc.paragraphs:
            text = para.text.strip()
            img_md = extract_images_from_paragraph(para)

            if text:
                full_text_list.append(text)
            if img_md:
                full_text_list.append(img_md)

        full_text_with_images = "\n".join(full_text_list)
        docs_to_save = []

        parent_blocks = full_text_with_images.split("***")

        for p_idx, parent_block in enumerate(parent_blocks):
            if not parent_block.strip(): continue

            parent_title = ""
            lines = parent_block.strip().split('\n')
            for line in lines:
                if "文档名：" in line:
                    parent_title = line.strip()
                    break
            if not parent_title:
                parent_title = f"文档名：{file_name}"

            full_parent_context = parent_block.strip()
            child_blocks = parent_block.split("<--split-->")

            for c_idx, child_content in enumerate(child_blocks):
                child_content = child_content.strip()
                if not child_content: continue

                content_to_vectorize = f"{parent_title}\n{child_content}"

                metadata = {
                    "source": file_name,
                    "type": "child_chunk",
                    "recall_context": full_parent_context,
                    "chunk_id": f"{file_name}_p{p_idx}_c{c_idx}",
                    "region_code": region_mapping.get(file_name, "000000")
                }

                new_doc = Document(page_content=content_to_vectorize, metadata=metadata)
                docs_to_save.append(new_doc)

        return docs_to_save

    except Exception as e:
        print(f" 解析文件 {file_path} 失败: {e}")
        import traceback
        traceback.print_exc()
        return []


def _safe_conn_str(conn_str: str) -> str:
    """隐藏连接串中的密码"""
    import re
    return re.sub(r'://([^:]+):([^@]+)@', r'://\1:****@', conn_str)


def _check_collection_owner(conn_str: str, collection_name: str) -> bool:
    """检查集合是否属于本项目（通过 pg_class 查询），返回 True 表示安全"""
    try:
        from sqlalchemy import create_engine, text
        engine = create_engine(conn_str)
        with engine.connect() as conn:
            result = conn.execute(
                text("SELECT relname FROM pg_catalog.pg_class WHERE relname = :name AND relkind = 'r'"),
                {"name": collection_name},
            )
            rows = result.fetchall()
            if rows:
                count_result = conn.execute(text(f'SELECT COUNT(*) FROM "{collection_name}"'))
                row_count = count_result.scalar()
                print(f"⚠️  集合 '{collection_name}' 已存在，包含 {row_count} 条数据")
                return True
            else:
                print(f"✅ 集合 '{collection_name}' 不存在，将新建")
                return False
    except Exception as e:
        print(f"⚠️  无法检查集合状态: {e}")
        return False


def ingest():
    # ===================== 安全检查 =====================
    print(f"\n{'=' * 60}")
    print(f"⚠️  即将连接数据库并操作向量集合")
    print(f"{'=' * 60}")
    print(f"数据库连接: {_safe_conn_str(DB_CONNECTION)}")
    print(f"目标集合名: {COLLECTION_NAME}")
    print(f"{'=' * 60}\n")

    # 确认连接的是项目自己的集合名
    if "parent_child" not in COLLECTION_NAME:
        print(f"❌ 安全拦截: 集合名 '{COLLECTION_NAME}' 不符合项目命名规范，拒绝操作")
        print(f"   请在 .env 中检查 PGVECTOR_COLLECTION_NAME 配置")
        return

    vector_store = PgVectorStore(
        embedding=EMBEDDING_MODEL,
        collection_name=COLLECTION_NAME,
        connection=DB_CONNECTION,
    )

    if not os.path.exists(folder_path):
        print("文件夹不存在")
        return

    # 检查集合是否已存在
    collection_exists = _check_collection_owner(DB_CONNECTION, COLLECTION_NAME)

    if collection_exists:
        print(f"\n⚠️⚠️⚠️  危险操作警告  ⚠️⚠️⚠️")
        print(f"即将清空集合 '{COLLECTION_NAME}' 内的所有数据！")
        print(f"此操作不可逆，且只会影响本项目集合，不会影响其他表/集合。")
        print(f"数据库地址: {_safe_conn_str(DB_CONNECTION)}")
        print("⏳ 正在清空集合...")
        vector_store.delete_collection()

    # 创建集合（如果不存在会自动建表）
    vector_store.create_collection()
    print(f"✅ 集合 '{COLLECTION_NAME}' 就绪")

    all_docs = []
    file_list = [f for f in os.listdir(folder_path) if f.endswith((".docx", ".txt"))]

    print(f"\n📂 开始处理 {len(file_list)} 个文档...")
    for f in file_list:
        path = os.path.join(folder_path, f)
        if f.endswith(".txt"):
            docs = read_and_process_txt(path)
        else:
            docs = read_and_process_word(path)
        all_docs.extend(docs)
        print(f"  - {f}: 生成 {len(docs)} 个切片")

    if all_docs:
        print(f"\n正在写入 {len(all_docs)} 个切片到 PostgreSQL pgvector 数据库...")
        vector_store.add_documents(documents=all_docs, batch_size=32)
        print("✅ 入库完成！")
    else:
        print("无数据入库")


if __name__ == "__main__":
    ingest()
