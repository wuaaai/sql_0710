from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(r"D:\pythonPro\dataQuery")
SOURCE_MD = ROOT / "dataQuery_技术架构设计文档.md"
FLOW_PNG = ROOT / "dataquery_flow_formal.png"
DOCX_PATH = ROOT / "dataquery_architecture_formal.docx"
PDF_PATH = ROOT / "dataquery_architecture_formal.pdf"


def load_source_markdown() -> str:
    """读取技术架构 Markdown 文档。"""
    return SOURCE_MD.read_text(encoding="utf-8")


def load_chinese_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    """加载本机可用的中文字体，用于绘制流程图。"""
    font_candidates = [
        (r"C:\Windows\Fonts\msyhbd.ttc", True),
        (r"C:\Windows\Fonts\msyh.ttc", False),
        (r"C:\Windows\Fonts\simsun.ttc", False),
    ]
    if bold:
        ordered = [font_candidates[0], font_candidates[1], font_candidates[2]]
    else:
        ordered = [font_candidates[1], font_candidates[2], font_candidates[0]]

    for font_path, _ in ordered:
        if Path(font_path).exists():
            return ImageFont.truetype(font_path, size=size)
    return ImageFont.load_default()


def wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.ImageFont,
    max_width: int,
) -> list[str]:
    """按指定宽度对流程图文字进行简单换行。"""
    lines: list[str] = []
    current = ""
    for char in text:
        test = current + char
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines or [text]


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    desc: str,
    fill: tuple[int, int, int],
    outline: tuple[int, int, int],
) -> None:
    """绘制流程图中的矩形节点。"""
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)

    title_font = load_chinese_font(26, bold=True)
    text_font = load_chinese_font(20)

    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    title_width = title_bbox[2] - title_bbox[0]
    draw.text(
        (x1 + (x2 - x1 - title_width) / 2, y1 + 18),
        title,
        font=title_font,
        fill=(34, 52, 85),
    )

    lines = wrap_text(draw, desc, text_font, max_width=(x2 - x1 - 36))
    current_y = y1 + 64
    for line in lines:
        line_bbox = draw.textbbox((0, 0), line, font=text_font)
        line_width = line_bbox[2] - line_bbox[0]
        draw.text(
            (x1 + (x2 - x1 - line_width) / 2, current_y),
            line,
            font=text_font,
            fill=(78, 96, 128),
        )
        current_y += 28


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: tuple[int, int, int] = (87, 112, 230),
) -> None:
    """绘制带箭头的连接线。"""
    draw.line([start, end], fill=color, width=4)
    arrow_size = 10
    ex, ey = end
    sx, sy = start

    if abs(ex - sx) > abs(ey - sy):
        if ex >= sx:
            points = [(ex, ey), (ex - arrow_size, ey - 6), (ex - arrow_size, ey + 6)]
        else:
            points = [(ex, ey), (ex + arrow_size, ey - 6), (ex + arrow_size, ey + 6)]
    else:
        if ey >= sy:
            points = [(ex, ey), (ex - 6, ey - arrow_size), (ex + 6, ey - arrow_size)]
        else:
            points = [(ex, ey), (ex - 6, ey + arrow_size), (ex + 6, ey + arrow_size)]
    draw.polygon(points, fill=color)


def generate_flow_image() -> Path:
    """生成“从用户输入问题到模型回复”的正式流程图。"""
    width, height = 1800, 2100
    image = Image.new("RGB", (width, height), (245, 248, 253))
    draw = ImageDraw.Draw(image)

    title_font = load_chinese_font(36, bold=True)
    sub_font = load_chinese_font(22)

    draw.text(
        (425, 40),
        "dataQuery 用户输入问题到模型回复流程图",
        font=title_font,
        fill=(22, 44, 88),
    )
    draw.text(
        (640, 96),
        "正式版技术架构设计文档配图",
        font=sub_font,
        fill=(92, 108, 132),
    )

    boxes = {
        "user": (520, 160, 1280, 300),
        "entry": (520, 370, 1280, 530),
        "route": (520, 610, 1280, 790),
        "qa": (120, 900, 820, 1090),
        "sql": (980, 900, 1680, 1090),
        "compose": (520, 1210, 1280, 1370),
        "output": (520, 1460, 1280, 1620),
        "final": (520, 1710, 1280, 1870),
    }

    draw_box(
        draw,
        boxes["user"],
        "1. 用户输入问题",
        "用户或 Dify 工作流传入自然语言问题和上下文参数。",
        (233, 242, 255),
        (126, 159, 255),
    )
    draw_box(
        draw,
        boxes["entry"],
        "2. 统一接入层",
        "server.py 解析 messages、region_code、stream 等入参，并进入统一编排服务。",
        (232, 247, 241),
        (93, 181, 140),
    )
    draw_box(
        draw,
        boxes["route"],
        "3. 路由编排层",
        "unified_qa_service 调用 router_service，完成主意图判断、子意图拆分和执行计划编排。",
        (255, 245, 229),
        (230, 163, 76),
    )
    draw_box(
        draw,
        boxes["qa"],
        "4A. 智能问答链路",
        "问题缺少明确科目、指标、收支方向、地区层级时，进入 RAG 文档检索与答案生成。",
        (245, 238, 255),
        (162, 126, 230),
    )
    draw_box(
        draw,
        boxes["sql"],
        "4B. 智能问数链路",
        "提取到明确业务语义后，进入业务域路由、表能力过滤、向量召回、SQL 编译与执行。",
        (255, 238, 240),
        (232, 116, 136),
    )
    draw_box(
        draw,
        boxes["compose"],
        "5. 结果整合层",
        "answer_composer 汇总问答结果、问数结果、摘要、表格明细和图表配置。",
        (233, 242, 255),
        (126, 159, 255),
    )
    draw_box(
        draw,
        boxes["output"],
        "6. 扩展输出层",
        "如可画图则缓存 option，/v1/chart/export 输出 ECharts option 供 Dify 模板渲染。",
        (232, 247, 241),
        (93, 181, 140),
    )
    draw_box(
        draw,
        boxes["final"],
        "7. 模型最终回复",
        "统一返回自然语言结论、数据明细、图表信息或文档依据，形成最终响应。",
        (255, 245, 229),
        (230, 163, 76),
    )

    draw_arrow(draw, (900, 300), (900, 370))
    draw_arrow(draw, (900, 530), (900, 610))
    draw_arrow(draw, (700, 790), (470, 900))
    draw_arrow(draw, (1100, 790), (1330, 900))
    draw_arrow(draw, (470, 1090), (820, 1210))
    draw_arrow(draw, (1330, 1090), (980, 1210))
    draw_arrow(draw, (900, 1370), (900, 1460))
    draw_arrow(draw, (900, 1620), (900, 1710))

    note_font = load_chinese_font(18)
    notes = [
        "路由关键规则：只有同时提取到明确科目、明确指标、明确收支方向、明确地区层级，才进入智能问数。",
        "若问题是复合句，则先拆分子意图，再分别执行智能问答或智能问数，最后合并回复。",
        "动态管理平台可更新文档向量库和智能问数表向量库，以支撑后续业务扩容。",
    ]
    y = 1930
    for note in notes:
        draw.rounded_rectangle(
            (140, y - 8, 1660, y + 28),
            radius=10,
            fill=(255, 255, 255),
            outline=(214, 223, 238),
            width=1,
        )
        draw.text((160, y), note, font=note_font, fill=(89, 101, 123))
        y += 44

    image.save(FLOW_PNG)
    return FLOW_PNG


def extract_headings(markdown_text: str) -> list[tuple[int, str]]:
    """提取 Markdown 中的二级标题，生成精简目录。"""
    headings: list[tuple[int, str]] = []
    for line in markdown_text.splitlines():
        if line.startswith("## "):
            headings.append((2, line[3:].strip()))
    return headings


def clean_inline_markdown(text: str) -> str:
    """清理行内 Markdown 标记，方便输出到 Word 或 PDF。"""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text.strip()


def add_page_number(section) -> None:
    """为 Word 页脚添加页码。"""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_word_flowchart(doc: Document, flow_image_path: Path) -> None:
    """在 Word 正文中插入流程图。"""
    doc.add_picture(str(flow_image_path), width=Cm(16.5))
    caption = doc.add_paragraph("图 1  dataQuery 从统一接入到最终回复的核心流程")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_word_document(markdown_text: str, flow_image_path: Path) -> Path:
    """生成正式版 Word 文档。"""
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(11)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("dataQuery 技术架构设计文档")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("正式版")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(74, 95, 139)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    meta = [
        "项目名称：财政智能问答与智能问数一体化系统",
        "文档范围：text_smart_qa / fiscal_smart_qa / 统一入口与动态向量库设计",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    ]
    for item in meta:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(item)

    doc.add_page_break()
    doc.add_heading("目录", level=1)
    for level, title in extract_headings(markdown_text):
        paragraph = doc.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Cm(0.8 if level == 2 else 1.6)
        paragraph.add_run(title)

    doc.add_page_break()

    in_code_block = False
    skip_mermaid = False

    for line in markdown_text.splitlines():
        stripped = line.strip()

        if stripped.startswith("```"):
            if stripped == "```mermaid":
                skip_mermaid = True
                add_word_flowchart(doc, flow_image_path)
            else:
                in_code_block = not in_code_block
            continue

        if skip_mermaid:
            if stripped == "```":
                skip_mermaid = False
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(clean_inline_markdown(line[3:]), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(clean_inline_markdown(line[4:]), level=2)
            continue
        if stripped == "---":
            doc.add_paragraph("")
            continue
        if re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(clean_inline_markdown(stripped), style="List Number")
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(clean_inline_markdown(stripped[2:]), style="List Bullet")
            continue
        if in_code_block:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(stripped)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
            continue

        paragraph = doc.add_paragraph(clean_inline_markdown(stripped))
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.3)
        add_page_number(section)

    doc.save(DOCX_PATH)
    return DOCX_PATH


def register_pdf_fonts() -> str:
    """注册 PDF 使用的中文字体。"""
    font_candidates = [
        ("MicrosoftYaHei", r"C:\Windows\Fonts\msyh.ttc"),
        ("SimSun", r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for name, path in font_candidates:
        if Path(path).exists():
            pdfmetrics.registerFont(TTFont(name, path))
            return name
    return "Helvetica"


def append_pdf_flowchart(story: list, flow_image_path: Path, styles) -> None:
    """在 PDF 正文中插入流程图。"""
    story.append(Spacer(1, 0.15 * cm))
    story.append(PdfImage(str(flow_image_path), width=17.2 * cm, height=20.1 * cm))
    story.append(Spacer(1, 0.15 * cm))
    story.append(Paragraph("图 1  dataQuery 从统一接入到最终回复的核心流程", styles["CnSubTitle"]))
    story.append(Spacer(1, 0.15 * cm))


def build_pdf_document(markdown_text: str, flow_image_path: Path) -> Path:
    """生成正式版 PDF 文档。"""
    font_name = register_pdf_fonts()
    styles = getSampleStyleSheet()

    styles.add(
        ParagraphStyle(
            name="CnTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=22,
            leading=30,
            textColor=colors.HexColor("#1d3557"),
            alignment=TA_CENTER,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CnSubTitle",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=14,
            leading=20,
            textColor=colors.HexColor("#4a5f8b"),
            alignment=TA_CENTER,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CnHeading1",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=16,
            leading=22,
            spaceBefore=12,
            spaceAfter=8,
            textColor=colors.HexColor("#1d3557"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="CnHeading2",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=13,
            leading=20,
            spaceBefore=10,
            spaceAfter=6,
            textColor=colors.HexColor("#314d79"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="CnBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=18,
            alignment=TA_JUSTIFY,
            firstLineIndent=21,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CnList",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=18,
            leftIndent=18,
            bulletIndent=6,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CnCode",
            parent=styles["Code"],
            fontName=font_name,
            fontSize=9.5,
            leading=14,
            backColor=colors.HexColor("#f4f6fb"),
            borderPadding=5,
        )
    )

    story: list = []
    story.append(Spacer(1, 3 * cm))
    story.append(Paragraph("dataQuery 技术架构设计文档", styles["CnTitle"]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("正式版", styles["CnSubTitle"]))
    story.append(Spacer(1, 0.8 * cm))
    story.append(Paragraph("项目名称：财政智能问答与智能问数一体化系统", styles["CnSubTitle"]))
    story.append(Paragraph("文档范围：text_smart_qa / fiscal_smart_qa / 统一入口与动态向量库设计", styles["CnSubTitle"]))
    story.append(Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles["CnSubTitle"]))
    story.append(PageBreak())

    story.append(Paragraph("目录", styles["CnHeading1"]))
    for level, title in extract_headings(markdown_text):
        prefix = "&nbsp;" * (4 if level == 3 else 0)
        story.append(Paragraph(f"{prefix}{title}", styles["CnBody"]))

    story.append(PageBreak())

    in_code_block = False
    skip_mermaid = False

    for line in markdown_text.splitlines():
        stripped = line.strip()

        if stripped.startswith("```"):
            if stripped == "```mermaid":
                skip_mermaid = True
                append_pdf_flowchart(story, flow_image_path, styles)
            else:
                in_code_block = not in_code_block
            continue

        if skip_mermaid:
            if stripped == "```":
                skip_mermaid = False
            continue

        if not stripped:
            story.append(Spacer(1, 0.15 * cm))
            continue

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            story.append(Paragraph(clean_inline_markdown(line[3:]), styles["CnHeading1"]))
            continue
        if line.startswith("### "):
            story.append(Paragraph(clean_inline_markdown(line[4:]), styles["CnHeading2"]))
            continue
        if stripped == "---":
            story.append(Spacer(1, 0.18 * cm))
            continue
        if re.match(r"^\d+\.\s+", stripped):
            story.append(Paragraph(clean_inline_markdown(stripped), styles["CnList"]))
            continue
        if stripped.startswith("- "):
            story.append(Paragraph(f"• {clean_inline_markdown(stripped[2:])}", styles["CnList"]))
            continue
        if in_code_block:
            story.append(Paragraph(clean_inline_markdown(stripped), styles["CnCode"]))
            continue

        story.append(Paragraph(clean_inline_markdown(stripped), styles["CnBody"]))

    def draw_page_number(canvas_obj, doc_obj) -> None:
        canvas_obj.setFont(font_name, 9)
        canvas_obj.setFillColor(colors.HexColor("#5f6c85"))
        canvas_obj.drawCentredString(A4[0] / 2, 1.1 * cm, f"{canvas_obj.getPageNumber()}")

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        topMargin=2.2 * cm,
        bottomMargin=1.8 * cm,
        leftMargin=2.1 * cm,
        rightMargin=2.0 * cm,
    )
    doc.build(story, onFirstPage=draw_page_number, onLaterPages=draw_page_number)
    return PDF_PATH


def main() -> None:
    """执行 Word 和 PDF 文档生成。"""
    markdown_text = load_source_markdown()
    flow_image_path = generate_flow_image()
    docx_file = build_word_document(markdown_text, flow_image_path)
    pdf_file = build_pdf_document(markdown_text, flow_image_path)
    print(f"WORD={docx_file}")
    print(f"PDF={pdf_file}")
    print(f"FLOW={flow_image_path}")


if __name__ == "__main__":
    main()
